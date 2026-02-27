# NOTE: This version skips execution if streamlit is not available.
try:
    import streamlit as st
    STREAMLIT_AVAILABLE = True
except ModuleNotFoundError:
    print("streamlit is not installed. Some features will be disabled.")
    STREAMLIT_AVAILABLE = False

import pdfplumber
from PIL import Image, ImageDraw, ImageFont
import pytesseract
import google.generativeai as genai
from docx import Document
import io
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

api_key = "AIzaSyCHSIpWiX5eHnElQo4074SmG3_4dhEF9gQ"

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def sanitize_text(text):
    return ''.join(c for c in text if c.isprintable())

def wrap_text(text, max_width, font_size):
    wrapped_lines = []
    lines = text.split('\n')
    for line in lines:
        while len(line) > max_width:
            split_index = line.rfind(' ', 0, max_width)
            if split_index == -1:
                split_index = max_width
            wrapped_lines.append(line[:split_index])
            line = line[split_index:].strip()
        wrapped_lines.append(line)
    return wrapped_lines

if STREAMLIT_AVAILABLE and api_key:
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-pro")
    chat = model.start_chat(history=[])

    st.title("Document Conversion and Chatbot")
    theme = st.radio("Choose Theme:", ["Light", "Dark"], horizontal=True)

    if theme == "Dark":
        st.markdown("""
            <style>
                body, .stApp {
                    background-color: #0e1117;
                    color: white;
                }
                input, textarea {
                    background-color: #1e1e1e !important;
                    color: white !important;
                }
            </style>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
            <style>
                body, .stApp {
                    background-color: white;
                    color: black;
                }
                input, textarea {
                    background-color: #ffffff !important;
                    color: black !important;
                }
            </style>
        """, unsafe_allow_html=True)

    uploaded_file = st.file_uploader("Upload an image, PDF, text, or Word file", type=["jpg", "jpeg", "png", "pdf", "txt", "docx"])

    if uploaded_file:
        filename = st.text_input("Enter the filename (without extension):", "output")
        text, sanitized_text = "", ""

        if uploaded_file.type == "application/pdf":
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
            st.subheader("Extracted Text from PDF")
            st.write(text)

            with open(f"{filename}.txt", "w", encoding="utf-8") as text_file:
                text_file.write(text)
            st.download_button(f"Download Extracted Text as {filename}.txt", text, file_name=f"{filename}.txt")

            doc = Document()
            sanitized_text = sanitize_text(text)
            doc.add_paragraph(sanitized_text)
            doc.save(f"{filename}.docx")
            with open(f"{filename}.docx", "rb") as doc_file:
                st.download_button(f"Download as Word Document ({filename}.docx)", doc_file, file_name=f"{filename}.docx")

        elif uploaded_file.type == "text/plain":
            text = uploaded_file.read().decode("utf-8")
            st.subheader("Uploaded Text")
            st.write(text)

            sanitized_text = sanitize_text(text)
            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            text_object = c.beginText(40, 750)
            text_object.setFont("Helvetica", 12)
            for line in wrap_text(sanitized_text, 100, 12):
                text_object.textLine(line)
            c.drawText(text_object)
            c.showPage()
            c.save()
            buffer.seek(0)
            st.download_button(f"Download Text as PDF ({filename}.pdf)", buffer, file_name=f"{filename}.pdf")

            img = Image.new('RGB', (800, 400), color='white')
            draw = ImageDraw.Draw(img)
            draw.text((10, 10), sanitized_text, font=ImageFont.load_default(), fill='black')
            img_buf = io.BytesIO()
            img.save(img_buf, format='JPEG')
            st.download_button(f"Download Text as JPEG ({filename}.jpeg)", img_buf.getvalue(), file_name=f"{filename}.jpeg")

            doc = Document()
            doc.add_paragraph(sanitized_text)
            doc.save(f"{filename}.docx")
            with open(f"{filename}.docx", "rb") as doc_file:
                st.download_button(f"Download Text as Word Document ({filename}.docx)", doc_file, file_name=f"{filename}.docx")

        elif uploaded_file.type in ["image/jpeg", "image/png"]:
            image = Image.open(uploaded_file)
            text = pytesseract.image_to_string(image)
            st.subheader("Extracted Text from Image")
            st.write(text)
            sanitized_text = sanitize_text(text)

            with open(f"{filename}.txt", "w", encoding="utf-8") as text_file:
                text_file.write(sanitized_text)
            st.download_button(f"Download Extracted Text as {filename}.txt", sanitized_text, file_name=f"{filename}.txt")

            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            text_object = c.beginText(40, 750)
            for line in wrap_text(sanitized_text, 100, 12):
                text_object.textLine(line)
            c.drawText(text_object)
            c.showPage()
            c.save()
            buffer.seek(0)
            st.download_button(f"Download Image Text as PDF ({filename}.pdf)", buffer, file_name=f"{filename}.pdf")

        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(uploaded_file)
            text = "\n".join([para.text for para in doc.paragraphs])
            st.subheader("Extracted Text from Word Document")
            st.write(text)

            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            text_object = c.beginText(40, 750)
            for line in wrap_text(text, 100, 12):
                text_object.textLine(line)
            c.drawText(text_object)
            c.showPage()
            c.save()
            buffer.seek(0)
            st.download_button(f"Download Word Text as PDF ({filename}.pdf)", buffer, file_name=f"{filename}.pdf")

        st.subheader("Ask a question based on the extracted text")
        user_question = st.text_input("Your Question:")

        if user_question and (text or sanitized_text):
            try:
                prompt = f"The following text was extracted from the document:\n\n{text or sanitized_text}\n\nBased on this, the user asks: {user_question}\n\nPlease provide an answer."
                response = chat.send_message(prompt)
                st.write(f"**Bot:** {response.text}")

                buffer = io.BytesIO()
                c = canvas.Canvas(buffer, pagesize=letter)
                text_object = c.beginText(40, 750)
                text_object.setFont("Helvetica", 12)
                text_object.textLines(["Chatbot Response:"] + wrap_text(response.text, 100, 12))
                c.drawText(text_object)
                c.showPage()
                c.save()
                buffer.seek(0)
                st.download_button(f"Download Chatbot Response as PDF ({filename}_response.pdf)", buffer, file_name=f"{filename}_response.pdf")
            except Exception as e:
                st.write(f"Error occurred: {e}")
else:
    print("streamlit is required to run the app. Please install it using 'pip install streamlit'")

